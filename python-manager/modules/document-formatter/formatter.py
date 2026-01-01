"""OnlyOffice-first formatter with cleanup to remove blank pages."""

import os
import time
from typing import Dict, Tuple

import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Standard formatting settings
STANDARD_FONT = "Times New Roman"
STANDARD_SIZE = 12  # pts
STANDARD_ALIGNMENT = WD_ALIGN_PARAGRAPH.JUSTIFY
STANDARD_LINE_SPACING = 1.15
STANDARD_MARGIN = Inches(1)  # 1 inch margins


def _paragraph_has_content(paragraph) -> bool:
    ns = paragraph._element.nsmap
    if paragraph.text and paragraph.text.strip():
        return True
    # Keep paragraphs that hold drawings/images
    return bool(paragraph._element.xpath(".//w:drawing", namespaces=ns))


def _remove_page_break_runs(paragraph) -> int:
    ns = paragraph._element.nsmap
    removed = 0
    for run in list(paragraph.runs):
        br_elems = run._element.xpath(".//w:br", namespaces=ns)
        if br_elems:
            parent = run._element.getparent()
            parent.remove(run._element)
            removed += 1
    return removed


def _apply_standard_formatting(doc: Document) -> Tuple[int, int, int, int]:
    paragraphs_formatted = 0
    tables_processed = 0
    total_runs_formatted = 0
    page_breaks_removed = 0

    for section in doc.sections:
        section.top_margin = STANDARD_MARGIN
        section.bottom_margin = STANDARD_MARGIN
        section.left_margin = STANDARD_MARGIN
        section.right_margin = STANDARD_MARGIN

    # Format paragraphs and strip empty/page-break-only paragraphs
    for paragraph in list(doc.paragraphs):
        page_breaks_removed += _remove_page_break_runs(paragraph)
        paragraph.alignment = STANDARD_ALIGNMENT
        fmt = paragraph.paragraph_format
        fmt.line_spacing = STANDARD_LINE_SPACING
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)

        for run in paragraph.runs:
            run.font.name = STANDARD_FONT
            run.font.size = Pt(STANDARD_SIZE)
            total_runs_formatted += 1

        if not _paragraph_has_content(paragraph):
            paragraph._element.getparent().remove(paragraph._element)
            continue

        paragraphs_formatted += 1

    # Format tables
    for table in doc.tables:
        tables_processed += 1
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _remove_page_break_runs(paragraph)
                    paragraph.alignment = STANDARD_ALIGNMENT
                    fmt = paragraph.paragraph_format
                    fmt.line_spacing = STANDARD_LINE_SPACING
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = STANDARD_FONT
                        run.font.size = Pt(STANDARD_SIZE)
                        total_runs_formatted += 1

    return paragraphs_formatted, tables_processed, total_runs_formatted, page_breaks_removed


def _format_with_python_docx(input_path: str, output_path: str) -> Dict[str, int]:
    doc = Document(input_path)
    p_cnt, t_cnt, r_cnt, br_cnt = _apply_standard_formatting(doc)
    doc.save(output_path)
    return {
        "paragraphs_formatted": p_cnt,
        "tables_processed": t_cnt,
        "total_runs_formatted": r_cnt,
        "page_break_runs_removed": br_cnt,
        "engine": "python-docx"
    }


def format_docx_via_onlyoffice(input_path: str, output_path: str) -> Dict[str, int]:
    server = os.getenv("ONLYOFFICE_SERVER", "http://localhost:8080").rstrip("/")
    convert_url = f"{server}/ConvertService.ashx"
    stats: Dict[str, int] = {}
    started = time.time()

    try:
        print(f"[formatter] Sending to OnlyOffice: {convert_url}")
        with open(input_path, "rb") as f:
            files = {
                "file": (
                    os.path.basename(input_path),
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            data = {
                "async": "false",
                "filetype": "docx",
                "outputtype": "docx",
                "title": os.path.basename(input_path),
            }
            resp = requests.post(convert_url, files=files, data=data, timeout=60)
            resp.raise_for_status()
            payload = resp.json()
            file_url = payload.get("fileUrl")
            if not file_url:
                raise RuntimeError("ConvertService response missing fileUrl")

        print(f"[formatter] OnlyOffice responded with fileUrl: {file_url}")
        download = requests.get(file_url, timeout=60)
        download.raise_for_status()
        with open(output_path, "wb") as out:
            out.write(download.content)

        # Post-process to remove blank paragraphs and page breaks
        stats = _format_with_python_docx(output_path, output_path)
        stats["engine"] = "onlyoffice+python-docx"
    except Exception as e:
        # Fallback to local formatting only
        print(f"[formatter] OnlyOffice failed, falling back to python-docx: {e}")
        stats = _format_with_python_docx(input_path, output_path)
        stats["engine"] = "python-docx-fallback"
        stats["fallback_error"] = str(e)

    stats["processing_time_ms"] = int((time.time() - started) * 1000)
    return stats


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Apply standard formatting to DOCX files")
    parser.add_argument("--input", required=True, help="Input DOCX file")
    parser.add_argument("--output", required=True, help="Output DOCX file")

    args = parser.parse_args()

    stats = format_docx_via_onlyoffice(args.input, args.output)

    print("Formatting complete:")
    print(f"  Engine: {stats.get('engine')}")
    print(f"  Paragraphs formatted: {stats['paragraphs_formatted']}")
    print(f"  Tables processed: {stats['tables_processed']}")
    print(f"  Total runs formatted: {stats['total_runs_formatted']}")
    print(f"  Page-break runs removed: {stats.get('page_break_runs_removed', 0)}")
    print(f"  Processing time: {stats['processing_time_ms']} ms")
    if stats.get("fallback_error"):
        print(f"  Fallback reason: {stats['fallback_error']}")
