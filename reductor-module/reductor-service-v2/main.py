"""
main.py

Reductor Service v2: FastAPI entry point
"""

import os
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from typing import Optional

from config import config
from logger import get_logger
from utils.minio_utils import minio_client
from utils.converter_utils import pdf_to_docx
from utils.identity_detector import detect_identity
from utils.docx_anonymizer import anonymize_docx, load_xml
from utils.docx_sanitizer import sanitize_docx_inplace
from pipeline.redact_pipeline import RedactionPipeline

logger = get_logger(__name__)

app = FastAPI(
    title=config.APP_NAME,
    version=config.APP_VERSION,
    debug=config.DEBUG,
)


# Models
class AnonymizeRequest(BaseModel):
    bucket: str
    object_key: str
    output_key: Optional[str] = None


class AnonymizeResponse(BaseModel):
    status: str
    anonymized_docx_path: Optional[str] = None
    minio_output_key: Optional[str] = None
    detected_before: Optional[dict] = None
    detected_after: Optional[dict] = None
    removed_bytes: int = 0


class PresidioRedactRequest(BaseModel):
    bucket: str
    object_key: str
    output_key: Optional[str] = None


class PresidioRedactResponse(BaseModel):
    status: str
    minio_output_key: str
    detections_count: int
    pii_types: list


# Routes
@app.get("/health")
def health():
    return {
        "status": "healthy",
        "service": "reductor-v2",
        "version": config.APP_VERSION,
    }


@app.post("/anonymize", response_model=AnonymizeResponse)
def anonymize(req: AnonymizeRequest):
    """
    End-to-end anonymization pipeline:
    1. Download PDF from MinIO raw/
    2. Convert PDF → DOCX
    3. Detect student identity
    4. Remove name + roll
    5. Upload anonymized DOCX to formatted/
    6. Return before/after report
    """
    try:
        logger.info(f"\n{'='*60}")
        logger.info(f"🚀 Starting anonymization pipeline")
        logger.info(f"   Bucket: {req.bucket}")
        logger.info(f"   Object: {req.object_key}")
        logger.info(f"{'='*60}")

        # Temp directory for local artifacts
        temp_dir = config.TEMP_DIR
        os.makedirs(temp_dir, exist_ok=True)
        
        filename_base = os.path.splitext(os.path.basename(req.object_key))[0]
        converted_path = os.path.join(temp_dir, f"{filename_base}_converted.docx")
        anonymized_path = os.path.join(temp_dir, f"{filename_base}_anonymized.docx")

        # Step 1: Download PDF from MinIO
        logger.info("\n[1/6] Downloading PDF from MinIO...")
        pdf_data = minio_client.download(req.bucket, req.object_key)

        # Step 2: Convert PDF → DOCX
        logger.info("\n[2/6] Converting PDF to DOCX...")
        docx_data = pdf_to_docx(pdf_data)
        with open(converted_path, "wb") as f:
            f.write(docx_data.getvalue())
        logger.info(f"✅ Saved to {converted_path}")

        # Sanitize converted DOCX to fix glyph artifacts
        logger.info("[2.5/6] Sanitizing converted DOCX for glyph cleanup...")
        sanitize_docx_inplace(converted_path)

        # Step 3: Detect identity BEFORE
        logger.info("\n[3/6] Detecting student identity (BEFORE anonymization)...")
        from utils.docx_anonymizer import unzip_docx
        import shutil
        temp_unzip = unzip_docx(converted_path)
        tree_before = load_xml(os.path.join(temp_unzip, "word/document.xml"))
        identity_before = detect_identity(tree_before)
        logger.info(f"✅ Detected: {identity_before}")
        shutil.rmtree(temp_unzip)

        # Step 4: Anonymize
        logger.info("\n[4/6] Anonymizing (removing name and roll)...")
        anon_stats = anonymize_docx(
            converted_path,
            anonymized_path,
            name=identity_before.get("name"),
            roll_no=identity_before.get("roll_no"),
        )

        # Step 5: Detect identity AFTER
        logger.info("\n[5/6] Detecting student identity (AFTER anonymization)...")
        temp_unzip = unzip_docx(anonymized_path)
        tree_after = load_xml(os.path.join(temp_unzip, "word/document.xml"))
        identity_after = detect_identity(tree_after)
        logger.info(f"✅ After anonymization: {identity_after}")
        shutil.rmtree(temp_unzip)

        # Step 6: Upload to MinIO
        logger.info("\n[6/6] Uploading anonymized DOCX to MinIO...")
        output_key = req.output_key or req.object_key.replace("/raw/", "/formatted/").replace(".pdf", "_anonymized.docx")
        with open(anonymized_path, "rb") as f:
            import io
            buf = io.BytesIO(f.read())
        minio_client.upload(
            req.bucket,
            output_key,
            buf,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        logger.info(f"\n{'='*60}")
        logger.info(f"✅ Anonymization pipeline COMPLETE")
        logger.info(f"{'='*60}\n")

        return AnonymizeResponse(
            status="success",
            anonymized_docx_path=anonymized_path,
            minio_output_key=output_key,
            detected_before=identity_before,
            detected_after=identity_after,
            removed_bytes=anon_stats.get("bytes_removed", 0),
        )

    except Exception as e:
        logger.error(f"\n❌ Anonymization FAILED: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/presidio-redact", response_model=PresidioRedactResponse)
def presidio_redact(req: PresidioRedactRequest):
    """
    Presidio-based PII redaction pipeline:
    1. Download PDF from MinIO
    2. Convert PDF → DOCX
    3. Extract text from DOCX
    4. Run Presidio + Regex PII detection
    5. Redact all detected PII entities (preserves labels like "Name:", "Roll No:")
    6. Save redacted DOCX back to MinIO
    7. Return detection report
    """
    try:
        logger.info(f"\n{'='*60}")
        logger.info(f"🔒 Starting Presidio PII redaction pipeline")
        logger.info(f"   Bucket: {req.bucket}")
        logger.info(f"   Object: {req.object_key}")
        logger.info(f"{'='*60}")

        # Initialize redaction pipeline
        pipeline = RedactionPipeline()

        # Temp directory for local artifacts
        temp_dir = config.TEMP_DIR
        os.makedirs(temp_dir, exist_ok=True)
        
        filename_base = os.path.splitext(os.path.basename(req.object_key))[0]
        converted_path = os.path.join(temp_dir, f"{filename_base}_converted.docx")
        redacted_path = os.path.join(temp_dir, f"{filename_base}_redacted.docx")

        # Step 1: Download PDF from MinIO
        logger.info("\n[1/5] Downloading PDF from MinIO...")
        pdf_data = minio_client.download(req.bucket, req.object_key)

        # Step 2: Convert PDF → DOCX
        logger.info("\n[2/5] Converting PDF to DOCX...")
        docx_data = pdf_to_docx(pdf_data)
        with open(converted_path, "wb") as f:
            f.write(docx_data.getvalue())
        logger.info(f"✅ Saved to {converted_path}")

        # Sanitize converted DOCX to fix glyph artifacts
        logger.info("[2.5/5] Sanitizing converted DOCX for glyph cleanup...")
        sanitize_docx_inplace(converted_path)

        # Step 3: Extract text from DOCX
        logger.info("\n[3/5] Extracting text from DOCX...")
        from docx import Document
        doc = Document(converted_path)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        logger.info(f"✅ Extracted {len(full_text)} characters")

        # Step 4: Run Presidio + Regex detection
        logger.info("\n[4/5] Running Presidio + Regex PII detection...")
        redacted_text, stats = pipeline.redact_text(full_text)
        
        # Get detection statistics from stats
        all_detections = stats.get('entities', [])
        pii_types = list(set([det['entity_type'] for det in all_detections]))
        logger.info(f"✅ Detected {len(all_detections)} PII entities: {pii_types}")

        # Step 5: Apply redactions to DOCX while preserving formatting
        logger.info("\n[5/5] Applying redactions to DOCX (preserving formatting & labels)...")
        
        # Use redaction pipeline directly on DOCX paragraph content
        doc = Document(converted_path)
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue
            
            # Get original paragraph text
            original_para_text = paragraph.text
            
            # Run through redaction pipeline (which preserves labels)
            redacted_para_text, _ = pipeline.redact_text(original_para_text)
            
            # If text was redacted, intelligently replace while preserving run formatting
            if original_para_text != redacted_para_text:
                # Build character position map from runs
                runs = list(paragraph.runs)
                if not runs:
                    continue
                
                # Calculate which characters belong to which run
                run_char_ranges = []
                char_pos = 0
                for run in runs:
                    run_len = len(run.text)
                    run_char_ranges.append((char_pos, char_pos + run_len, run))
                    char_pos += run_len
                
                # Distribute redacted text across existing runs sequentially, preserving formatting
                # This keeps bullets/numbering alignment intact.
                red_idx = 0
                total_len = len(redacted_para_text)
                for run in runs:
                    orig_len = len(run.text)
                    if orig_len <= 0:
                        continue
                    end_idx = min(red_idx + orig_len, total_len)
                    segment = redacted_para_text[red_idx:end_idx]
                    run.text = segment
                    red_idx = end_idx
                # Append any remaining characters to the last run without clearing styles
                if red_idx < total_len:
                    runs[-1].text = (runs[-1].text or "") + redacted_para_text[red_idx:]
                
                logger.debug(f"Redacted paragraph: {len(original_para_text)} → {len(redacted_para_text)} chars")
        
        doc.save(redacted_path)
        logger.info(f"✅ Redactions applied to {redacted_path} (labels preserved)")

        # Step 6: Upload to MinIO

        # Step 6: Upload to MinIO
        logger.info("\n[6/6] Uploading redacted DOCX to MinIO...")
        output_key = req.output_key or req.object_key.replace("/raw/", "/anonymized/").replace(".pdf", "_redacted.docx")
        with open(redacted_path, "rb") as f:
            import io
            buf = io.BytesIO(f.read())
        minio_client.upload(
            req.bucket,
            output_key,
            buf,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        logger.info(f"\n{'='*60}")
        logger.info(f"✅ Presidio redaction pipeline COMPLETE")
        logger.info(f"{'='*60}\n")

        return PresidioRedactResponse(
            status="success",
            minio_output_key=output_key,
            detections_count=len(all_detections),
            pii_types=pii_types,
        )

    except Exception as e:
        logger.error(f"\n❌ Presidio redaction FAILED: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=config.PORT,
        log_level="info",
    )
